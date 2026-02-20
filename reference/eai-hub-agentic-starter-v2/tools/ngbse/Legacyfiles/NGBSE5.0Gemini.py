# setup_ngbse5.py
import os
import shutil
import textwrap
from pathlib import Path

print(">>> Bouwen van NGBSE 5.0: Fusion Engine <<<")
root = Path("./NGBSE-5.0-FUSION")
if root.exists(): shutil.rmtree(root)
print(f"Projectstructuur wordt aangemaakt in: {root.resolve()}")
for d in ["engine", "collectors", "report", "config", "data", "out", "enrich", "export"]:
    (root / d).mkdir(parents=True, exist_ok=True)

# --- CONFIGURATIE ---
(root / "config/ngbse.config.yml").write_text(textwrap.dedent('''
collectors:
  enabled: ["censys", "leakix", "shodan", "wayback", "urlscan", "github"]
orchestrator:
  inter_collector_sleep_seconds: 0.2
  max_pages_per_collector: 2
  use_cache: true
  cache_dir: "cache"
  cache_ttl_seconds: 3600
  logfile: "out/run.log.jsonl"
scoring:
  # Kwaliteitsweging per bron (Q-factor)
  quality_weights:
    shodan: 1.0; censys: 1.0; leakix: 0.9; github: 0.8; urlscan: 0.7; wayback: 0.5; default: 0.4
  # E_AI* modelgewichten
  eai_weights: { w1: 0.3, w2: 0.3, w3: 0.2, w4: 0.15, w5: 0.05 }
  # Standaardwaarden voor E_AI* parameters (P en V zijn dynamisch)
  eai_defaults:
    W_V: 0.7; D_A: 0.7; D_B: 0.7; T: 0.7; A: 0.7
exports:
  csv: true
  stix21: true
'''))
(root / "data/seeds.jsonl").write_text(textwrap.dedent('''
{"seed": "blob.core.windows.net?sv=", "where": "urlscan", "why": "Azure SAS token sporen"}
{"seed": "token.actions.githubusercontent.com aud", "where": "github", "why": "OIDC trust policies"}
{"seed": "port:1883 MQTT country:NL", "where": "shodan", "why": "Open MQTT brokers in NL"}
{"seed": "allow_anonymous true", "where": "leakix", "why": "Anonieme toegang-configuraties"}
{"seed": "services.software.vendor:\\"Citrix\\" AND country:NL", "where": "censys", "why": "Citrix ADC/NetScaler aanwezigheid"}
{"seed": "*citrix.com/*/release-notes/*", "where": "wayback", "why": "Mirror-aanwezigheid van oude configs"}
'''))
(root / ".env.example").write_text("# Vul uw API-sleutels in en hernoem naar .env\nURLSCAN_API_KEY=\nGITHUB_TOKEN=\nSHODAN_API_KEY=\nCENSYS_API_ID=\nCENSYS_API_SECRET=\nLEAKIX_API_KEY=\n")
(root / "requirements.txt").write_text("requests\nPyYAML\npython-docx\npython-dotenv\n")

# --- KERN ENGINE ---
(root / "engine/orchestrator.py").write_text(textwrap.dedent('''
import yaml, json, time, argparse, hashlib
from pathlib import Path
from dotenv import load_dotenv
from engine.registry import load_collectors
from engine.fuzzy_dedupe import dedupe
from enrich.enrichment import run_enrichment
from engine.scoring import run_scoring
from enrich.cluster import cluster_by_domain
from enrich.timeline import timeline_by_day
from report.make_report import build_report
from export.csv_export import write_csv
from export.stix21 import write_stix_bundle

def main():
    parser = argparse.ArgumentParser(description="NGBSE 5.0: Fusion Engine")
    args = parser.parse_args() # Simplified for autorun
    
    cfg = yaml.safe_load(Path("config/ngbse.config.yml").read_text())
    load_dotenv()

    # 1. VERZAMELEN
    print("[1/5] Verzamelen van data (incl. CoC Hashing)...")
    collectors = load_collectors(cfg)
    seeds = [json.loads(line) for line in Path("data/seeds.jsonl").read_text().splitlines()]
    raw_findings = []
    for seed in seeds:
        collector = collectors.get(seed['where'])
        if collector:
            for finding in collector.collect(seed): raw_findings.append(finding)
            time.sleep(cfg['orchestrator']['inter_collector_sleep_seconds'])
    
    # 2. ONTDUBBELEN
    print(f"[2/5] Ontdubbelen van {len(raw_findings)} ruwe bevindingen...")
    unique_findings = dedupe(raw_findings)
    
    # 3. VERRIJKEN
    print(f"[3/5] Verrijken van {len(unique_findings)} unieke bevindingen (M, C, Q)...")
    enriched_findings = run_enrichment(unique_findings, cfg)

    # 4. SCOREN
    print("[4/5] Toepassen van strategische E_AI* scoring...")
    scored_findings = run_scoring(enriched_findings, cfg)

    # 5. OPSLAAN & RAPPORTEREN
    print("[5/5] Genereren van outputs en rapporten...")
    out_dir = Path("out"); out_dir.mkdir(exist_ok=True)
    findings_path = out_dir / "findings.jsonl"
    with findings_path.open("w", encoding="utf-8") as f:
        for finding in scored_findings: f.write(json.dumps(finding) + "\\n")
    
    clusters = cluster_by_domain(scored_findings)
    (out_dir / "clusters.json").write_text(json.dumps(clusters, indent=2))
    
    timeline = timeline_by_day(scored_findings)
    (out_dir / "timeline.json").write_text(json.dumps(timeline, indent=2))

    if cfg.get("exports", {}).get("csv"):
        write_csv(scored_findings, str(out_dir / "findings.csv"))
    if cfg.get("exports", {}).get("stix21"):
        write_stix_bundle(scored_findings, str(out_dir / "findings.stix.json"))
        
    report_path = build_report(scored_findings, clusters, timeline)
    print(f"\\n>>> NGBSE 5.0 Run Voltooid. Strategisch Rapport opgeslagen in: {report_path} <<<")

if __name__ == "__main__": main()
'''))
(root / "engine/registry.py").write_text(textwrap.dedent('''
import importlib
def load_collectors(config):
    mods = {}
    for name in config.get("collectors", {}).get("enabled", []):
        try: mods[name] = importlib.import_module(f"collectors.{name}_collector")
        except ImportError: print(f"Kon collector '{name}' niet laden.")
    return mods
'''))
(root / "engine/fuzzy_dedupe.py").write_text(textwrap.dedent('''
import re, hashlib
from urllib.parse import urlparse
def normalize_url(u):
    if not u or not isinstance(u, str): return ""
    p = urlparse(u.lower()); host = p.hostname or ""; host = host[4:] if host.startswith("www.") else host
    return f"{host}{p.path or '/'}"
def soft_hash(f):
    key = f"{f['seed'].get('seed', '')}|{f['source']}|{normalize_url(f.get('url'))}"
    return hashlib.sha256(key.encode()).hexdigest()
def dedupe(findings):
    seen=set(); out=[]
    for f in findings:
        h = soft_hash(f)
        if h not in seen: seen.add(h); out.append(f)
    return out
'''))
(root / "engine/scoring.py").write_text(textwrap.dedent('''
import math
def e_ai_star(params, weights):
    # V' = Validity Prime, de gewogen validiteit inclusief Kwaliteit en Corroboratie
    Vp = min(1.0, 0.5 * params['V'] + 0.3 * params['Q'] + 0.2 * min(1.0, params['C'] / 3.0))
    score = math.sqrt(
        weights['w1'] * (params['P'] * params['W_V']) + weights['w2'] * (params['D_A'] * params['D_B']) +
        weights['w3'] * (params['T'] * params['A']) + weights['w4'] * Vp + weights['w5'] * params['M']
    )
    return round(score, 4)
def run_scoring(findings, config):
    cfg = config.get('scoring', {})
    weights = cfg.get('eai_weights', {})
    defaults = cfg.get('eai_defaults', {})
    for f in findings:
        enrich = f.get('enrichment', {})
        params = {
            'P': 1.0, 'V': 0.5, # Aanwezigheid is 1.0, Validiteit start op 0.5 (niet geverifieerd)
            'Q': enrich.get('Q', 0.4), 'C': enrich.get('C', 1), 'M': enrich.get('M', 0.2), **defaults
        }
        f['score'] = e_ai_star(params, weights)
    findings.sort(key=lambda x: x.get('score', 0), reverse=True)
    return findings
'''))

# --- VERRIJKING (ENRICHMENT) ---
(root / "enrich/enrichment.py").write_text(textwrap.dedent('''
from collections import defaultdict
def run_enrichment(findings, config):
    by_seed = defaultdict(list)
    for f in findings: by_seed[f['seed']['seed']].append(f)
    q_weights = config.get('scoring', {}).get('quality_weights', {})
    for seed_str, items in by_seed.items():
        live = {i['source'] for i in items if i['source'] != 'wayback'}
        C = len(live)
        M = 1.0 if 'wayback' in {i['source'] for i in items} and not live else 0.6 if 'wayback' in {i['source'] for i in items} else 0.2
        for item in items:
            Q = q_weights.get(item['source'], q_weights.get('default', 0.4))
            item['enrichment'] = {'M': M, 'C': C, 'Q': Q}
    return findings
'''))
(root / "enrich/cluster.py").write_text(textwrap.dedent('''
from urllib.parse import urlparse
from collections import defaultdict
def _host(u):
    try: return urlparse(u).hostname or ""
    except: return ""
def cluster_by_domain(findings):
    clusters=defaultdict(list)
    for f in findings:
        host = _host(f.get("url") or f.get("where", ""))
        if host: clusters[host].append({"source": f["source"], "score": f["score"]})
    return {k: v for k, v in sorted(clusters.items(), key=lambda item: len(item[1]), reverse=True)}
'''))
(root / "enrich/timeline.py").write_text(textwrap.dedent('''
from collections import defaultdict
def timeline_by_day(findings):
    by=defaultdict(int)
    for f in findings:
        if ts := f.get("ts"): by[str(ts)[:10]] += 1
    return dict(sorted(by.items()))
'''))

# --- COLLECTORS (met CoC) ---
(root / "collectors/util.py").write_text("import hashlib\ndef coc_sha256(material: str) -> str: return hashlib.sha256(material.encode('utf-8', 'ignore')).hexdigest()")
# Voorbeeld voor urlscan, de rest volgt hetzelfde patroon
(root / "collectors/urlscan_collector.py").write_text(textwrap.dedent('''
import os, requests
from collectors.util import coc_sha256
def collect(seed):
    q = seed.get("seed"); api_key = os.getenv("URLSCAN_API_KEY")
    headers = {"API-Key": api_key} if api_key else {}
    params = {"q": q, "size": 50}
    try:
        r = requests.get("https://urlscan.io/api/v1/search/", params=params, headers=headers, timeout=20)
        r.raise_for_status()
        for item in r.json().get("results", []):
            url=item.get("page", {}).get("url"); ts=item.get("indexedAt"); rid=item.get("_id")
            material = f"{rid}|{url}|{ts}"
            yield {"source": "urlscan", "seed": seed, "url": url, "ts": ts, "coc_sha256": coc_sha256(material)}
    except requests.RequestException as e:
        print(f"URLScan Error: {e}")
        yield from []
'''))
# Placeholders voor andere collectors
for name in ["github", "shodan", "censys", "leakix", "wayback"]:
    (root / f"collectors/{name}_collector.py").write_text(f"def collect(seed): print(f'Skipping {name} collector...'); yield from []")

# --- EXPORTS ---
(root / "export/csv_export.py").write_text("import csv\ndef write_csv(findings, path):\n    if not findings: return\n    keys=findings[0].keys()\n    with open(path,'w',newline='',encoding='utf-8') as f:\n        w=csv.DictWriter(f,fieldnames=keys);w.writeheader();w.writerows(findings)")
(root / "export/stix21.py").write_text(textwrap.dedent('''
import json, uuid, datetime
def _id(tp): return f"{tp}--{uuid.uuid4()}"
def write_stix_bundle(findings, path):
    objs=[]
    for f in findings[:200]:
        if not (url := f.get("url")): continue
        objs.append({"type":"indicator","spec_version":"2.1","id":_id("indicator"), "pattern_type":"stix",
                     "pattern":f"[url:value = '{url}']", "name": f"OSINT Finding: {url}",
                     "created":datetime.datetime.utcnow().isoformat()+"Z",
                     "valid_from":(f.get('ts') or datetime.datetime.utcnow().isoformat()+'Z')})
    with open(path,"w") as fo: json.dump({"type":"bundle","id":_id("bundle"),"objects":objs}, fo, indent=2)
'''))

# --- RAPPORTAGE ---
(root / "report/make_report.py").write_text(textwrap.dedent('''
import json
from pathlib import Path
from docx import Document
def build_report(findings, clusters, timeline):
    out_path = Path("out/NGBSE-5.0_Fusion_Report.docx")
    doc = Document()
    doc.add_heading("NGBSE 5.0: Fusion Engine Report", 0)
    doc.add_paragraph(f"Totaal unieke, strategisch gescoorde bevindingen: {len(findings)}")
    doc.add_heading("Top Clusters (op domein)", 1)
    for domain, items in list(clusters.items())[:5]: doc.add_paragraph(f"- {domain} ({len(items)} hits)")
    doc.add_heading("Strategische Bevindingen (Top 20)", 1)
    table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
    hdr=table.rows[0].cells; hdr[0].text="E_AI*"; hdr[1].text="Info"; hdr[2].text="URL / Waar"; hdr[3].text="CoC Hash"
    for f in findings[:20]:
        row=table.add_row().cells
        row[0].text=f"{f.get('score', 0):.4f}"
        enrich=f.get('enrichment',{}); row[1].text=f"M={enrich.get('M')} C={enrich.get('C')} Q={enrich.get('Q')}"
        row[2].text=f.get('url', f.get('where',''))[:80]
        row[3].text=f.get('coc_sha256', '')[:16]
    doc.save(out_path)
    return out_path
'''))

# --- RUN SCRIPT ---
(root / "run.sh").write_text(textwrap.dedent('''
#!/bin/bash
set -e
echo ">>> NGBSE 5.0: Fusion Engine starten..."
if [ ! -f ".env" ]; then
    echo "INFO: .env niet gevonden, .env.example wordt gekopieerd."
    cp .env.example .env
    echo "BELANGRIJK: Vul uw API-sleutels in .env voor resultaten."
fi
pip install -q -r requirements.txt
python -m engine.orchestrator
'''))
os.chmod(root / "run.sh", 0o755)

print("\\n>>> SETUP VOLTOOID. NGBSE 5.0 is gebouwd. <<<")
print(f"1. Navigeer naar de map: cd {root}")
print("2. Vul uw API-sleutels in het '.env' bestand.")
print("3. Voer de engine uit met: ./run.sh")