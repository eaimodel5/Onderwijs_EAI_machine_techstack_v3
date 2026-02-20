import json, hashlib, time
from pathlib import Path
from docx import Document

BASE = Path(__file__).resolve().parents[1]
FINDINGS = BASE/"data/findings.jsonl"
TRIAGED  = BASE/"data/triaged.jsonl"
OUTDOC   = BASE/"NL_Blindspot_Engine_v2_Report.docx"

def sha256_file(p:Path)->str:
    h=hashlib.sha256()
    with p.open("rb") as f:
        h.update(f.read())
    return h.hexdigest()

def main():
    doc = Document()
    doc.add_heading("NL Blindspot Engine v2 — Report", 0)
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%SZ', time.gmtime())}")
    doc.add_paragraph("Scope: clearnet presence-only. No content extraction or auth required.")
    doc.add_heading("Findings (presence)", level=1)
    if FINDINGS.exists():
        with FINDINGS.open("r", encoding="utf-8") as f:
            for i, line in enumerate(f):
                try:
                    j = json.loads(line)
                except:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"[{i+1}] ").bold=True
                p.add_run(f"{j.get('source')}: {j.get('page')} | q={j.get('query')} | ts={j.get('ts')}")
    else:
        doc.add_paragraph("No findings yet.")

    doc.add_heading("Mirror Triaged", level=1)
    if TRIAGED.exists():
        with TRIAGED.open("r", encoding="utf-8") as f:
            for i, line in enumerate(f):
                j = json.loads(line)
                p = doc.add_paragraph()
                p.add_run(f"[M{i+1}] ").bold=True
                p.add_run(f"{j.get('page')} | M={j.get('mirror_factor')} | ts={j.get('ts_triage')}")
    else:
        doc.add_paragraph("No triaged mirrors.")

    doc.add_heading("Chain-of-Custody", level=1)
    if FINDINGS.exists():
        doc.add_paragraph(f"findings.jsonl SHA-256: {sha256_file(FINDINGS)}")
    if TRIAGED.exists():
        doc.add_paragraph(f"triaged.jsonl  SHA-256: {sha256_file(TRIAGED)}")

    doc.save(OUTDOC)
    print(str(OUTDOC))

if __name__ == "__main__":
    main()
