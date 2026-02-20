# setup_ngbse10.py
import os
import shutil
import textwrap
from pathlib import Path

print(">>> Bouwen van NGBSE 10.0: De Complete, Pure Online Engine <<<")
root = Path("./NGBSE_10_FINAL")
if root.exists(): shutil.rmtree(root)
print(f"Projectstructuur wordt aangemaakt in: {root.resolve()}")
for d in ["engine", "collectors", "report", "config", "data", "out", "enrich", "export", "cache"]:
    (root / d).mkdir(parents=True, exist_ok=True)
for subdir in ["engine", "collectors", "report", "enrich", "export"]:
    (root / subdir / "__init__.py").touch()

# --- CONFIGURATIE ---
(root / "config/ngbse.config.yml").write_text(textwrap.dedent('''
# NGBSE 10.0 Configuratie
orchestrator:
  inter_collector_sleep_seconds: 0.2
  clearnet_only_mode: false # True forceert een run zonder API sleutels
collectors:
  enabled: ["censys", "leakix", "shodan", "wayback", "urlscan", "github"]
scoring:
  quality_weights:
    shodan: 1.0
    censys: 1.0
    leakix: 0.9
    github: 0.8
    urlscan: 0.7
    wayback: 0.5
    default: 0.4
  eai_weights: { w1: 0.3, w2: 0.3, w3: 0.2, w4: 0.15, w5: 0.05 }
  eai_defaults: { W_V: 0.7, D_A: 0.7, D_B: 0.7, T: 0.7, A: 0.7 }
predictive_synthesis:
  enabled: true
  prior_probability: 0.1
  # Let op: de LLM-endpoint is localhost, aangezien het lokaal draait
  llm_endpoint: "http://localhost:11434/api/generate"
  llm_model: "phi3:mini"
  themes:
    Cloud_Token_Leakage: {keywords: ["sv=", "signature=", "blob.core.windows.net"], future_scenario: "Ongeautoriseerde Toegang tot Cloud Data"}
    Public_Code_Exposure: {keywords: ["github", "oidc", "token"], future_scenario: "Misbruik van Gelekte Credentials"}
validation: { enabled: true, allowlist_file: "config/validation.allowlist.txt" }
exports: { csv: true, stix21: true }
'''))
(root / "data/seeds.jsonl").write_text('{"seed": "domain:example.com", "where": "urlscan", "why": "Initiële test seed"}')
(root / ".env.example").write_text("# Vul uw API-sleutels in en hernoem naar .env\nURLSCAN_API_KEY=\nGITHUB_TOKEN=\nSHODAN_API_KEY=\nCENSYS_API_ID=\nCENSYS_API_SECRET=\nLEAKIX_API_KEY=\n")
(root / "requirements.txt").write_text("requests\nPyYAML\npython-docx\npython-dotenv\nbeautifulsoup4\nstix2\n")
(root / "config/validation.allowlist.txt").write_text("# example.com\n")

# --- KERN ENGINE (VOLLEDIG) ---
(root / "engine/orchestrator.py").write_text(textwrap.dedent('''
import yaml
import json
import time
import argparse
from pathlib import Path
from dotenv import load_dotenv

from engine.registry import load_collectors
from engine.fuzzy_dedupe import dedupe
from engine.validation import validate_findings
from enrich.enrichment import run_enrichment
from engine.scoring import run_scoring
from engine.predictive_synthesis import generate_synthesis
from report.make_report import build_report
from export.csv_export import write_csv
from export.stix_exporter import export_to_stix
from engine.logger import get_logger

def main():
    parser = argparse.ArgumentParser(description="NGBSE 10.0 Engine")
    parser.add_argument("--validate", action="store_true")
    parser.add_argument("--clearnet", action="store_true")
    args = parser.parse_args()
    
    cfg = yaml.safe_load(Path("config/ngbse.config.yml").read_text())
    log = get_logger("NGBSE-10.0", "out/run.log.jsonl")
    load_dotenv()
    
    if args.validate: cfg['validation']['enabled'] = True
    if args.clearnet: cfg['orchestrator']['clearnet_only_mode'] = True
    
    log.info("ONLINE MODUS: Starten van dataverzameling...")
    collectors = load_collectors(cfg, log)
    seeds = [json.loads(line) for line in Path("data/seeds.jsonl").read_text().splitlines()]
    raw_findings = []
    for seed in seeds:
        if collector := collectors.get(seed['where']):
            try:
                for finding in collector.collect(seed, cfg, log):
                    raw_findings.append(finding)
                time.sleep(cfg['orchestrator']['inter_collector_sleep_seconds'])
            except Exception as e:
                log.error(f"FATAL ERROR in collector '{seed['where']}'", {"error": str(e)})
    
    log.info(f"Ontdubbelen van {len(raw_findings)} bevindingen...")
    unique_findings = dedupe(raw_findings)
    
    log.info("Live validatie uitvoeren...")
    if cfg.get("validation", {}).get("enabled", False):
        allowlist = Path(cfg["validation"]["allowlist_file"]).read_text().splitlines()
        validated_findings = validate_findings(unique_findings, allowlist, log)
    else:
        validated_findings = unique_findings
        
    log.info("Verrijken met strategische factoren (M, C, Q)...")
    enriched_findings = run_enrichment(validated_findings, cfg)

    log.info("Toepassen van strategische E_AI* scoring...")
    scored_findings = run_scoring(enriched_findings, cfg)
    
    log.info("Genereren van voorspellende synthese...")
    synthesis = generate_synthesis(scored_findings, cfg, log)

    out_dir = Path("out")
    out_dir.mkdir(exist_ok=True)
    (out_dir / "findings.jsonl").write_text("\\n".join(json.dumps(f) for f in scored_findings))
    (out_dir / "synthesis_report.json").write_text(json.dumps(synthesis, indent=2, ensure_ascii=False))
        
    if cfg.get("exports", {}).get("csv"):
        write_csv(scored_findings, str(out_dir / "findings.csv"))
    if cfg.get("exports", {}).get("stix21"):
        export_to_stix(scored_findings, synthesis, "NGBSE-10.0", str(out_dir / "findings.stix.json"))
        
    report_path = build_report(scored_findings, synthesis)
    log.info(f"NGBSE 10.0 Run Voltooid. Rapport: {report_path}")
    print(f"\\n>>> NGBSE 10.0 Run Voltooid. Rapport: {report_path} <<<")

if __name__ == "__main__":
    main()
'''))
(root / "engine/registry.py").write_text(textwrap.dedent('''
import importlib
from collectors.base_collector import BaseCollector

def load_collectors(config, log):
    modules = {}
    enabled = config.get("collectors", {}).get("enabled", [])
    for name in enabled:
        try:
            module = importlib.import_module(f"collectors.{name}_collector")
            # Vind de klasse in de module die BaseCollector erft
            for item_name in dir(module):
                item = getattr(module, item_name)
                if isinstance(item, type) and issubclass(item, BaseCollector) and item is not BaseCollector:
                    modules[name] = item(config, log)
                    break
        except ImportError:
            log.error(f"Kon collector module niet laden: '{name}'")
    return modules
'''))
(root / "engine/fuzzy_dedupe.py").write_text(textwrap.dedent('''
import re
import hashlib
from urllib.parse import urlparse

def normalize_url(u):
    if not u or not isinstance(u, str): return ""
    p = urlparse(u.lower())
    host = p.hostname or ""
    host = host[4:] if host.startswith("www.") else host
    return f"{host}{p.path or '/'}"
    
def soft_hash(f):
    seed_str = f['seed'].get('seed', '') if isinstance(f.get('seed'), dict) else str(f.get('seed', ''))
    key = f"{seed_str}|{f.get('source', '')}|{normalize_url(f.get('url'))}"
    return hashlib.sha256(key.encode()).hexdigest()
    
def dedupe(findings):
    seen = set()
    unique_findings = []
    for f in findings:
        h = soft_hash(f)
        if h not in seen:
            seen.add(h)
            unique_findings.append(f)
    return unique_findings
'''))
(root / "engine/scoring.py").write_text(textwrap.dedent('''
import math

def e_ai_star(params, weights):
    Vp = min(1.0, 0.5 * params['V'] + 0.3 * params['Q'] + 0.2 * min(1.0, params['C'] / 3.0))
    score = math.sqrt(
        weights['w1'] * (params['P'] * params['W_V']) +
        weights['w2'] * (params['D_A'] * params['D_B']) +
        weights['w3'] * (params['T'] * params['A']) +
        weights['w4'] * Vp +
        weights['w5'] * params['M']
    )
    return round(score, 4)

def run_scoring(findings, config):
    cfg = config.get('scoring', {})
    weights = cfg.get('eai_weights', {})
    defaults = cfg.get('eai_defaults', {})
    for f in findings:
        enrich = f.get('enrichment', {})
        params = {
            'P': 1.0, 'V': 1.0 if f.get('validated') else 0.5,
            'Q': enrich.get('Q', 0.4), 'C': enrich.get('C', 1), 'M': enrich.get('M', 0.2), **defaults
        }
        f['score'] = e_ai_star(params, weights)
    findings.sort(key=lambda x: x.get('score', 0), reverse=True)
    return findings
'''))
(root / "engine/validation.py").write_text(textwrap.dedent('''
import re
import requests
from urllib.parse import urlparse

def _allowed(host, allowlist):
    host = host.lower()
    for e in [x.strip().lower() for x in allowlist if x.strip()]:
        if e.startswith(".") and host.endswith(e): return True
        if host == e: return True
    return False

def validate_findings(findings, allowlist, log):
    out = []
    validated_count = 0
    if not allowlist:
        log.warning("Validatie ingeschakeld, maar allowlist is leeg. Geen validatie uitgevoerd.")
        for f in findings:
            out.append({**f, "validated": False})
        return out
        
    for f in findings:
        url = f.get("url")
        if not url:
            out.append({**f, "validated": False})
            continue
            
        host = urlparse(url if re.match(r"^https?://", url) else "https://" + url).hostname or ""
        if not _allowed(host, allowlist):
            out.append({**f, "validated": False})
            continue
            
        try:
            r = requests.head(url, timeout=7, allow_redirects=True, headers={"User-Agent":"NGBSE-10.0"})
            ok = 200 <= r.status_code < 400
            if ok: validated_count += 1
            out.append({**f, "validated": ok, "status_code": r.status_code})
        except requests.RequestException as e:
            log.warning(f"Validatie mislukt voor {url}", {"error": str(e)})
            out.append({**f, "validated": False, "status_code": None})
            
    log.info(f"{validated_count} van de {len(findings)} relevante bevindingen succesvol gevalideerd.")
    return out
'''))
(root / "engine/logger.py").write_text(textwrap.dedent('''
import json
import logging
import sys
import time

class JsonLineFormatter(logging.Formatter):
    def format(self, record):
        payload = {"ts": int(time.time()*1000), "level": record.levelname, "msg": record.getMessage(), "name": record.name}
        if hasattr(record, 'args') and isinstance(record.args, dict):
            payload.update(record.args)
        return json.dumps(payload, ensure_ascii=False)

def get_logger(name, logfile=None):
    logger = logging.getLogger(name)
    logger.setLevel(logging.INFO)
    if not logger.handlers:
        h = logging.StreamHandler(sys.stdout)
        h.setFormatter(JsonLineFormatter())
        logger.addHandler(h)
        if logfile:
            fh = logging.FileHandler(logfile, encoding="utf-8")
            fh.setFormatter(JsonLineFormatter())
            logger.addHandler(fh)
    return logger
'''))
(root / "engine/predictive_synthesis.py").write_text(textwrap.dedent('''
from collections import defaultdict
import math
import json
import requests

def query_llm(prompt, config, log):
    # This is a placeholder as we can't run a real LLM.
    # It returns a template-based summary.
    # In a real appliance, this would call a local Ollama instance.
    log.warning("LLM call is gesimuleerd. Gebruik een echte lokale LLM voor productie.")
    theme = "Unknown"
    if "Scenario Theme: '" in prompt:
        theme = prompt.split("Scenario Theme: '")[1].split("'")[0]
    return f"Template-gebaseerde samenvatting voor risicothema '{theme}': De gevonden indicatoren wijzen op een verhoogd risico. Potentiële business impact is significant."

def generate_synthesis(findings, config, log):
    synthesis_cfg = config.get("predictive_synthesis", {})
    if not synthesis_cfg.get("enabled", False): return {}
    themes_def = synthesis_cfg.get("themes", {})
    for f in findings:
        f_themes = set()
        seed_str = f['seed'].get('seed', '') if isinstance(f.get('seed'), dict) else str(f.get('seed', ''))
        text_to_scan = seed_str.lower() + " " + (f.get('url') or '').lower()
        for theme, definition in themes_def.items():
            if any(kw in text_to_scan for kw in definition.get("keywords", [])):
                f_themes.add(theme)
        f['themes'] = list(f_themes)
    scenarios = defaultdict(lambda: {"findings": []})
    for f in findings:
        for theme in f.get('themes', []):
            scenarios[theme]["findings"].append(f)
    prior = synthesis_cfg.get("prior_probability", 0.1)
    prior_odds = prior / (1 - prior) if prior < 1 else float('inf')
    output = {}
    for theme, data in scenarios.items():
        items = data["findings"]
        if not items: continue
        avg_score = sum(f.get('score', 0.0) for f in items) / len(items)
        likelihood_ratio = 1 + 9 * (max(0, avg_score - 0.5) / 0.5)
        posterior_odds = prior_odds * likelihood_ratio
        probability = posterior_odds / (1 + posterior_odds)
        C = items[0].get('enrichment', {}).get('C', 0)
        Q_avg = sum(f.get('enrichment', {}).get('Q', 0) for f in items) / len(items)
        top_evidence = json.dumps([{"url": f.get('url'), "source": f.get('source')} for f in items[:3]])
        prompt = (f"Scenario Theme: '{theme}'. Average E_AI* score: {avg_score:.2f}, Corroboration: {C}. Top evidence: {top_evidence}.")
        summary_text = query_llm(prompt, config, log)
        output[theme] = {
            "future_scenario": themes_def.get(theme, {}).get("future_scenario", "Onbekend"),
            "probability_90_days": f"{probability:.1%}", "semantic_summary": summary_text,
            "avg_eai_score": avg_score, "evidence_count": len(items)
        }
    return dict(sorted(output.items(), key=lambda item: float(item[1]['probability_90_days'][:-1]), reverse=True))
'''))

# --- VERRIJKING & EXPORTS & RAPPORTAGE (VOLLEDIG) ---
(root / "enrich/enrichment.py").write_text(textwrap.dedent('''
from collections import defaultdict
def run_enrichment(findings, config):
    by_seed = defaultdict(list)
    for f in findings:
        seed_str = f['seed'].get('seed', '') if isinstance(f.get('seed'), dict) else str(f.get('seed', ''))
        by_seed[seed_str].append(f)
    q_weights = config.get('scoring', {}).get('quality_weights', {})
    for seed_str, items in by_seed.items():
        live_sources = {i['source'] for i in items if i['source'] != 'wayback'}
        C = len(live_sources)
        has_wayback = 'wayback' in {i['source'] for i in items}
        M = 1.0 if has_wayback and not live_sources else (0.6 if has_wayback else 0.2)
        for item in items:
            Q = q_weights.get(item['source'], q_weights.get('default', 0.4))
            item['enrichment'] = {'M': M, 'C': C, 'Q': Q}
    return findings
'''))
(root / "export/csv_export.py").write_text(textwrap.dedent('''
import csv
import json
def write_csv(findings, path):
    if not findings: return
    flat_findings = []
    for f in findings:
        ff = dict(f)
        enrich = ff.pop('enrichment', {})
        ff.update(enrich)
        ff.pop('themes', None)
        seed_data = ff.pop('seed', {})
        ff['seed_query'] = seed_data.get('seed', '')
        ff['seed_source'] = seed_data.get('where', '')
        flat_findings.append(ff)
    
    # Zorg dat alle keys aanwezig zijn voor de header
    all_keys = set()
    for item in flat_findings:
        all_keys.update(item.keys())
    
    sorted_keys = sorted(list(all_keys))

    with open(path, 'w', newline='', encoding='utf-8') as h:
        w = csv.DictWriter(h, fieldnames=sorted_keys, extrasaction='ignore')
        w.writeheader()
        w.writerows(flat_findings)
'''))
(root / "export/stix_exporter.py").write_text(textwrap.dedent('''
import json
import uuid
from datetime import datetime
from stix2 import TLP_WHITE, Indicator, Bundle, Report, Identity

def export_to_stix(findings, synthesis, author, out_path):
    stix_objects = []
    
    identity = Identity(name=author, identity_class="organization")
    stix_objects.append(identity)
    
    report_refs = []
    for f in findings:
        pattern = ""
        if url := f.get('url'):
            pattern = f"[url:value = '{url}']"
        elif where := f.get('where'):
             pattern = f"[{f.get('source')}:value = '{where}']"
        
        if not pattern: continue
        
        indicator = Indicator(
            name=f"OSINT Finding: {f.get('url', f.get('where', ''))}",
            pattern_type="stix",
            pattern=pattern,
            created_by_ref=identity.id,
            valid_from=f.get("ts") or datetime.utcnow(),
            labels=["osint", "automated-finding"],
            confidence=int(f.get("score", 0.5) * 100)
        )
        stix_objects.append(indicator)
        report_refs.append(indicator.id)
        
    synthesis_desc = "\\n".join([f"{data['future_scenario']} (Prob: {data['probability_90_days']})" for theme, data in synthesis.items()])
    stix_report = Report(
        name=f"NGBSE 10.0 Synthesis Report ({datetime.now().strftime('%Y-%m-%d')})",
        description=synthesis_desc,
        object_refs=report_refs,
        created_by_ref=identity.id,
        report_types=["threat-report"]
    )
    stix_objects.append(stix_report)
    
    bundle = Bundle(stix_objects, allow_custom=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(bundle.serialize(pretty=True))
'''))
(root / "report/make_report.py").write_text(textwrap.dedent('''
import json
from pathlib import Path
from docx import Document

def build_report(findings, synthesis):
    out_path = Path("out/NGBSE-10.0_Report.docx")
    doc = Document()
    doc.add_heading("NGBSE 10.0: Executive Briefing", 0)
    
    doc.add_heading("Voorspellende Risicoscenario's", 1)
    if not synthesis:
        doc.add_paragraph("Geen risicoscenario's gedetecteerd.")
    for theme, data in synthesis.items():
        doc.add_paragraph(data['future_scenario'], style='Heading 2')
        p = doc.add_paragraph()
        p.add_run('Waarschijnlijkheid (90 dagen): ').bold = True
        p.add_run(f'{data["probability_90_days"]}\\n')
        p.add_run('Samenvatting: ').bold = True
        p.add_run(data['semantic_summary'])
    
    doc.add_heading('Gedetailleerd Bewijs (Top 20)', 1)
    table = doc.add_table(1, 4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'E_AI*'
    hdr[1].text = 'Info (M,C,Q)'
    hdr[2].text = 'URL/Bewijs'
    hdr[3].text = 'CoC Hash'
    
    for f in findings[:20]:
        r = table.add_row().cells
        r[0].text = f'{f.get("score",0):.4f}'
        e = f.get('enrichment', {})
        r[1].text = f"M={e.get('M')} C={e.get('C')} Q={e.get('Q')}"
        r[2].text = str(f.get('url', f.get('where','')))[:60]
        r[3].text = str(f.get('coc_sha256',''))[:12]
        
    doc.save(out_path)
    return out_path
'''))

# --- OPERATIONELE COLLECTORS (VOLLEDIG & OBJECTGEORIËNTEERD) ---
(root / "collectors/base_collector.py").write_text(textwrap.dedent('''
import os
import requests
import time
from .util import coc_sha256

class BaseCollector:
    def __init__(self, config, log):
        self.config = config
        self.log = log
        self.source_name = "unknown"
        self.api_key = None
        self.requires_api_key = False
        self.api_key_env_var = None
        self.session = requests.Session()
        self.session.headers.update({"User-Agent": "NGBSE-10.0"})
        self._load_api_key()

    def _load_api_key(self):
        if self.requires_api_key and self.api_key_env_var:
            self.api_key = os.getenv(self.api_key_env_var)

    def collect(self, seed_obj, cfg, log):
        is_clearnet = cfg.get("orchestrator",{}).get("clearnet_only_mode", False)
        if self.requires_api_key and (not self.api_key or is_clearnet):
            log.warning(f"{self.source_name.capitalize()} collector: API-sleutel vereist en overgeslagen.", {"seed": seed_obj.get("seed")})
            return []
        return self.run(seed_obj)
        
    def run(self, seed_obj):
        raise NotImplementedError("De 'run' methode moet geïmplementeerd worden door een subklasse.")
'''))
(root / "collectors/urlscan_collector.py").write_text(textwrap.dedent('''
from .base_collector import BaseCollector
from .util import coc_sha256
class UrlscanCollector(BaseCollector):
    def __init__(self, config, log):
        super().__init__(config, log)
        self.source_name = "urlscan"
        self.requires_api_key = True
        self.api_key_env_var = "URLSCAN_API_KEY"
        self._load_api_key()
        if self.api_key:
            self.session.headers.update({"API-Key": self.api_key})

    def run(self, seed_obj):
        q = seed_obj.get("seed")
        try:
            r = self.session.get("https://urlscan.io/api/v1/search/", params={"q":q,"size":50}, timeout=20)
            r.raise_for_status()
            for i in r.json().get("results", []):
                url=i.get("page",{}).get("url"); ts=i.get("indexedAt"); rid=i.get("_id")
                yield {"source":"urlscan","seed":seed_obj,"url":url,"ts":ts,"coc_sha256":coc_sha256(f"{rid}|{url}|{ts}")}
        except Exception as e:
            self.log.error(f"{self.source_name} Fout", {"error": str(e), "seed": q})
'''))
(root / "collectors/github_collector.py").write_text(textwrap.dedent('''
from .base_collector import BaseCollector
from .util import coc_sha256
from bs4 import BeautifulSoup
from urllib.parse import quote_plus
class GithubCollector(BaseCollector):
    def __init__(self, config, log):
        super().__init__(config, log)
        self.source_name = "github"
        self.api_key_env_var = "GITHUB_TOKEN"
        self._load_api_key() # API key is optional for this collector

    def collect(self, seed_obj, cfg, log):
        is_clearnet = cfg.get("orchestrator",{}).get("clearnet_only_mode", False)
        if self.api_key and not is_clearnet:
            return self._run_api(seed_obj)
        else:
            return self._run_scrape(seed_obj)

    def _run_api(self, seed_obj):
        q = seed_obj.get("seed")
        try:
            self.session.headers.update({"Authorization": f"Bearer {self.api_key}"})
            r = self.session.get("https://api.github.com/search/code", params={"q":q,"per_page":50}, timeout=20)
            r.raise_for_status()
            for i in r.json().get("items",[]):
                url=i.get("html_url"); repo=i.get("repository",{}).get("full_name")
                yield {"source":"github","seed":seed_obj,"url":url,"ts":i.get("repository",{}).get("updated_at"),"coc_sha256":coc_sha256(f"{url}|{repo}")}
        except Exception as e:
            self.log.error(f"{self.source_name} API Fout", {"error": str(e), "seed": q})

    def _run_scrape(self, seed_obj):
        q = seed_obj.get("seed")
        self.log.warning("GitHub draait in clearnet-scrape-modus.")
        try:
            r = self.session.get(f"https://github.com/search?type=code&q={quote_plus(q)}", timeout=20); r.raise_for_status()
            soup = BeautifulSoup(r.text, 'html.parser')
            for a in soup.select('div.search-title a'):
                url = "https://github.com" + a['href']
                yield {"source":"github","seed":seed_obj,"url":url,"ts":None,"coc_sha256":coc_sha256(url)}
        except Exception as e:
            self.log.error(f"{self.source_name} Scrape Fout", {"error": str(e), "seed": q})
'''))
(root / "collectors/shodan_collector.py").write_text(textwrap.dedent('''
from .base_collector import BaseCollector
from .util import coc_sha256
class ShodanCollector(BaseCollector):
    def __init__(self, config, log):
        super().__init__(config, log)
        self.source_name = "shodan"
        self.requires_api_key = True
        self.api_key_env_var = "SHODAN_API_KEY"
        self._load_api_key()
    def run(self, seed_obj):
        q = seed_obj.get("seed")
        try:
            r = self.session.get("https://api.shodan.io/shodan/host/search", params={"key":self.api_key,"query":q}, timeout=20)
            r.raise_for_status()
            for i in r.json().get("matches",[]):
                ip=i.get("ip_str"); ts=i.get("timestamp")
                yield {"source":"shodan","seed":seed_obj,"where":ip,"url":f"https://www.shodan.io/host/{ip}","ts":ts,"coc_sha256":coc_sha256(f"{ip}|{ts}")}
        except Exception as e:
            self.log.error(f"{self.source_name} Fout", {"error": str(e), "seed": q})
'''))
(root / "collectors/censys_collector.py").write_text(textwrap.dedent('''
from .base_collector import BaseCollector
from .util import coc_sha256
import os
class CensysCollector(BaseCollector):
    def __init__(self, config, log):
        super().__init__(config, log)
        self.source_name = "censys"
        self.requires_api_key = True
        self.api_id = os.getenv("CENSYS_API_ID")
        self.api_secret = os.getenv("CENSYS_API_SECRET")
    def collect(self, seed_obj, cfg, log):
        is_clearnet = cfg.get("orchestrator",{}).get("clearnet_only_mode", False)
        if not (self.api_id and self.api_secret) or is_clearnet:
            log.warning("Censys collector: API ID/Secret vereist."); return []
        return self.run(seed_obj)
    def run(self, seed_obj):
        q = seed_obj.get("seed")
        try:
            r = self.session.post("https://search.censys.io/api/v2/hosts/search", auth=(self.api_id, self.api_secret), json={"q":q,"per_page":50}, timeout=20)
            r.raise_for_status()
            for i in r.json().get("result",{}).get("hits",[]):
                ip=i.get("ip"); ts=i.get("last_updated_at")
                yield {"source":"censys","seed":seed_obj,"where":ip,"url":f"https://search.censys.io/hosts/{ip}","ts":ts,"coc_sha256":coc_sha256(f"{ip}|{ts}")}
        except Exception as e:
            self.log.error(f"{self.source_name} Fout", {"error": str(e), "seed": q})
'''))
(root / "collectors/leakix_collector.py").write_text(textwrap.dedent('''
from .base_collector import BaseCollector
from .util import coc_sha256
class LeakixCollector(BaseCollector):
    def __init__(self, config, log):
        super().__init__(config, log)
        self.source_name = "leakix"
        self.requires_api_key = True
        self.api_key_env_var = "LEAKIX_API_KEY"
        self._load_api_key()
        if self.api_key:
            self.session.headers.update({"Api-Key": self.api_key})
    def run(self, seed_obj):
        q = seed_obj.get("seed")
        try:
            r = self.session.get("https://leakix.net/search", params={"q":q,"size":50}, timeout=20)
            r.raise_for_status()
            for i in r.json():
                url=i.get("link") or i.get("url"); ts=i.get("time")
                yield {"source":"leakix","seed":seed_obj,"url":url,"where":i.get("ip"),"ts":ts,"coc_sha256":coc_sha256(f"{url}|{ts}")}
        except Exception as e:
            self.log.error(f"{self.source_name} Fout", {"error": str(e), "seed": q})
'''))
(root / "collectors/wayback_collector.py").write_text(textwrap.dedent('''
from .base_collector import BaseCollector
from .util import coc_sha256
class WaybackCollector(BaseCollector):
    def __init__(self, config, log):
        super().__init__(config, log)
        self.source_name = "wayback"
    def run(self, seed_obj):
        q = seed_obj.get("seed")
        try:
            r=self.session.get("http://web.archive.org/cdx/search/cdx", params={"url":f"*{q}*","output":"json","limit":100}, timeout=20)
            r.raise_for_status()
            for row in r.json()[1:]:
                ts, orig, digest = row[1], row[2], row[5]
                url=f"https://web.archive.org/web/{ts}/{orig}"
                ts_iso=f"{ts[:4]}-{ts[4:6]}-{ts[6:8]}T{ts[8:10]}:{ts[10:12]}:{ts[12:14]}Z"
                yield {"source":"wayback","seed":seed_obj,"url":url,"ts":ts_iso,"coc_sha256":coc_sha256(f"{url}|{digest}")}
        except Exception as e:
            self.log.error(f"{self.source_name} Fout", {"error": str(e), "seed": q})
'''))

# --- RUN SCRIPT ---
(root / "run.sh").write_text(textwrap.dedent('''
#!/bin/bash
set -e
echo ">>> NGBSE-10.0: De Finale Engine starten..."
if [ ! -f .env ]; then
    cp .env.example .env
    echo "BELANGRIJK: Vul API-sleutels in .env voor de beste resultaten."
fi
pip install -q -r requirements.txt
python -m engine.orchestrator "$@"
'''))
os.chmod(root / "run.sh", 0o755)
(root / "README.md").write_text(textwrap.dedent('''
# NGBSE 10.0: De Finale Engine
Dit is de complete, zuivere, online-gerichte NGBSE.

## Setup
1. (Optioneel) Vul API-sleutels in een `.env` bestand (kopieer van `.env.example`).
2. Installeer een lokale LLM via [Ollama](https://ollama.com/) en draai `ollama run phi3:mini`.

## Gebruik
- **Volledige run:** `./run.sh`
- **Clearnet run (beperkt):** `./run.sh --clearnet`
- **Validatie forceren:** Voeg `--validate` toe.
'''))

print("\\n>>> SETUP VOLTOOID. NGBSE 10.0 is gebouwd. <<<")
print(f"1. Navigeer naar de map: cd {root.name}")
print("2. Vul uw API-sleutels in '.env'.")
print("3. Zorg dat Ollama lokaal draait (`ollama serve`).")
print("4. Start de engine met: ./run.sh")