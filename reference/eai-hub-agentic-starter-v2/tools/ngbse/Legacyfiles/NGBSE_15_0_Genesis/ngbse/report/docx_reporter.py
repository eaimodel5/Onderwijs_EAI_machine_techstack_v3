from docx import Document
from docx.shared import Pt
from typing import List, Dict, Any
from datetime import datetime

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def write_report(path: str, brief: Dict[str,Any], blindspots: Dict[str,Any], forecast: Dict[str,Any]):
    doc = Document()
    add_heading(doc, "NGBSE 15.0 – Intelligence Brief", 0)
    add_paragraph(doc, f"Generated: {datetime.utcnow().isoformat()}Z")

    add_heading(doc, "BLUF", 1)
    add_paragraph(doc, brief.get("BLUF",""))

    add_heading(doc, "Top Risks", 1)
    for r in brief.get("TopRisks", []):
        add_paragraph(doc, f"- {r['asset']}: risk {r['risk']}")

    add_heading(doc, "Actions", 1)
    for a in brief.get("Actions", []):
        add_paragraph(doc, f"- {a}")

    add_heading(doc, "Blindspot Analysis", 1)
    add_paragraph(doc, f"Coverage Ratios: {blindspots['coverage'].get('ratios',{})}")
    add_paragraph(doc, f"Coverage Imbalance: {blindspots['coverage'].get('imbalance',0.0)}")
    add_paragraph(doc, f"Recency Spike: {blindspots['recency'].get('spike',False)} / {blindspots['recency'].get('detail',{})}")
    add_paragraph(doc, f"Confidence Gap (risky assets): {blindspots['confidence'].get('risky_assets',[])}")

    add_heading(doc, "Forecast & Scenarios", 1)
    if forecast:
        for asset, fvals in forecast.items():
            add_paragraph(doc, f"- {asset}: last={round(fvals['last'],3)} ewma={round(fvals['ewma'],3)} → projected_next={round(fvals['projected_next'],3)}")
    else:
        add_paragraph(doc, "No historical series available yet.")

    add_heading(doc, "Analyst Review & Sign-off", 1)
    add_paragraph(doc, "Analyst Name: ________________________")
    add_paragraph(doc, "Date of Review: ______________________")
    add_paragraph(doc, "Conclusion: [Agree] [Disagree]")
    add_paragraph(doc, "Signature: ___________________________")

    doc.save(path)
