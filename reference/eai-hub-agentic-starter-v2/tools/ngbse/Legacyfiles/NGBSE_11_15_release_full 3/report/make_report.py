
import json
from pathlib import Path
from docx import Document

def build_report(findings, synthesis, blindspots):
    out = Path("out/NGBSE_11_Report.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    d = Document()
    d.add_heading("NGBSE 11.1: Final Genesis Report", 0)
    d.add_heading("Voorspellende Risicoscenario's", 1)
    if not synthesis: d.add_paragraph("Geen scenario's gedetecteerd.")
    for theme, data in synthesis.items():
        d.add_paragraph(theme, style="Heading 2")
        d.add_paragraph(f"Waarschijnlijkheid (90d): {data.get('probability_90_days')}")
        d.add_paragraph(f"Samenvatting: {data.get('semantic_summary','')}")
    d.add_heading("Gecombineerde Blindspots & Aanbevelingen", 1)
    for b in (blindspots or []):
        d.add_paragraph(b, style="List Bullet")
    d.add_heading("Bewijs (Top 20)", 1)
    table = d.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "E_AI*", "Bron", "URL", "CoC Hash"
    for f in findings[:20]:
        row = table.add_row().cells
        row[0].text = f"{f.get('score',0):.4f}"
        row[1].text = f.get("source","")
        row[2].text = (f.get("url") or f.get("where",""))[:80]
        row[3].text = (f.get("coc_sha256") or "")[:16]
    d.save(out)
    return out
