# setup_ngbse14_final_synthesis.py
import os
import shutil
import textwrap
from pathlib import Path

print(">>> Bouwen van NGBSE 14.0 'Final Synthesis': De Complete Online Engine <<<")
root = Path("./NGBSE_14_FINAL_SYNTHESIS")
if root.exists(): shutil.rmtree(root)
print(f"Projectstructuur wordt aangemaakt in: {root.resolve()}")
for d in ["engine", "collectors", "report", "config", "data", "out", "enrich", "export", "cache"]:
    (root / d).mkdir(parents=True, exist_ok=True)
for subdir in ["engine", "collectors", "report", "enrich", "export"]:
    (root / subdir / "__init__.py").touch()

# --- CONFIGURATIE (VOLLEDIG) ---
(root / "config/ngbse.config.yml").write_text(textwrap.dedent('''
# NGBSE 14.0 "Final Synthesis" Configuratie
orchestrator:
  inter_collector_sleep_seconds: 0.2
collectors:
  enabled: ["censys", "leakix", "shodan", "wayback", "urlscan", "github"]
scoring:
  quality_weights:
    shodan: 1.0; censys: 1.0; leakix: 0.9; github: 0.8; urlscan: 0.7; wayback: 0.5; default: 0.4
  eai_weights: { w1: 0.3, w2: 0.3, w3: 0.2, w4: 0.15, w5: 0.05 }
  eai_defaults: { W_V: 0.7, D_A: 0.7, D_B: 0.7, T: 0.7, A: 0.7 }
predictive_synthesis:
  enabled: true
  prior_probability: 0.1
  llm_provider: "google_gemini" # Pure online LLM
  llm_model: "gemini-1.5-flash"
  themes:
    Cloud_Token_Leakage: {keywords: ["sv=", "signature=", "blob.core.windows.net"], future_scenario: "Ongeautoriseerde Toegang tot Cloud Data"}
    Public_Code_Exposure: {keywords: ["github", "oidc", "token"], future_scenario: "Misbruik van Gelekte Credentials"}
reverse_llm:
  enabled: true
validation: { enabled: true, allowlist_file: "config/validation.allowlist.txt" }
exports: { csv: true, stix21: true }
'''))
(root / "data/seeds.jsonl").write_text('{"seed": "domain:example.com", "where": "urlscan", "why": "Initiële test seed"}')
(root / ".env.example").write_text("URLSCAN_API_KEY=\nGITHUB_TOKEN=\nSHODAN_API_KEY=\nCENSYS_API_ID=\nCENSYS_API_SECRET=\nLEAKIX_API_KEY=\n# VERPLICHT voor de online synthese\nGEMINI_API_KEY=\n")
(root / "requirements.txt").write_text("requests\nPyYAML\npython-docx\npython-dotenv\nbeautifulsoup4\nstix2\ngoogle-generativeai\n")
(root / "config/validation.allowlist.txt").write_text("# example.com\n")

# --- KERN ENGINE (VOLLEDIG & GECORRIGEERD) ---
(root / "engine/orchestrator.py").write_text(textwrap.dedent('''
import yaml, json, time, argparse
from pathlib import Path
from dotenv import load_dotenv
from engine.registry import load_collectors
from enrich.asset_clustering import cluster_by_asset
from engine.validation import validate_assets
from enrich.enrichment import run_enrichment
from engine.scoring import run_scoring
from engine.predictive_synthesis import generate_synthesis
from engine.reverse_llm import analyze_blindspots
from report.make_report import build_report
from export.csv_export import write_csv
from export.stix_exporter import export_to_stix
from engine.logger import get_logger

def main():
    cfg = yaml.safe_load(Path("config/ngbse.config.yml").read_text())
    log = get_logger("NGBSE-14.0", "out/run.log.jsonl")
    load_dotenv()
    
    log.info("ONLINE MODUS: Starten van dataverzameling...")
    collectors = load_collectors(cfg, log)
    seeds = [json.loads(line) for line in Path("data/seeds.jsonl").read_text().splitlines()]
    raw_findings = [f for s in seeds if (c := collectors.get(s['where'])) for f in c.collect(s, cfg, log)]
    
    log.info(f"Clusteren van {len(raw_findings)} bevindingen naar unieke assets...")
    asset_clusters = cluster_by_asset(raw_findings)
    
    log.info("Live validatie uitvoeren op unieke assets...")
    if cfg.get("validation", {}).get("enabled", False):
        allowlist = Path(cfg["validation"]["allowlist_file"]).read_text().splitlines()
        validated_clusters = validate_assets(asset_clusters, allowlist, log)
    else: validated_clusters = asset_clusters
        
    log.info("Verrijken van asset clusters (M, C, Q)...")
    enriched_clusters = run_enrichment(validated_clusters, cfg)

    log.info("Toepassen van dynamische, strategische E_AI* scoring op assets...")
    scored_assets = run_scoring(enriched_clusters, cfg)
    
    log.info("Genereren van voorspellende synthese met online LLM...")
    synthesis = generate_synthesis(scored_assets, cfg, log)

    log.info("Analyseren van gecombineerde blindspots...")
    blindspots = analyze_blindspots(scored_assets, synthesis)

    all_scored_findings = sorted([f for asset in scored_assets.values() for f in asset['findings']], key=lambda x: x.get('score', 0), reverse=True)
    
    out_dir = Path("out"); out_dir.mkdir(exist_ok=True)
    (out_dir / "findings.jsonl").write_text("\\n".join(json.dumps(f) for f in all_scored_findings))
    (out_dir / "synthesis_report.json").write_text(json.dumps(synthesis, indent=2, ensure_ascii=False))
    (out_dir / "blindspots.json").write_text(json.dumps(blindspots, indent=2, ensure_ascii=False))
        
    if cfg.get("exports", {}).get("csv"): write_csv(all_scored_findings, str(out_dir / "findings.csv"))
    if cfg.get("exports", {}).get("stix21"): export_to_stix(all_scored_findings, synthesis, "NGBSE-14.0", str(out_dir / "findings.stix.json"))
        
    report_path = build_report(scored_assets, synthesis, blindspots)
    log.info(f"NGBSE 14.0 Run Voltooid. Rapport: {report_path}")
    print(f"\\n>>> NGBSE 14.0 Run Voltooid. Rapport: {report_path} <<<")

if __name__ == "__main__": main()
'''))

(root / "engine/logger.py").write_text(textwrap.dedent('''
import json, logging, sys, time
class JsonLineFormatter(logging.Formatter):
    def format(self, record):
        payload = {"ts": int(time.time()*1000), "level": record.levelname, "msg": record.getMessage(), "name": record.name}
        if hasattr(record, 'args') and isinstance(record.args, dict): payload.update(record.args)
        return json.dumps(payload, ensure_ascii=False)
def get_logger(name, logfile=None):
    logger = logging.getLogger(name); logger.setLevel(logging.INFO)
    if not logger.handlers:
        h = logging.StreamHandler(sys.stdout); h.setFormatter(JsonLineFormatter()); logger.addHandler(h)
        if logfile:
            fh = logging.FileHandler(logfile, encoding="utf-8"); fh.setFormatter(JsonLineFormatter()); logger.addHandler(fh)
    return logger
'''))

(root / "engine/registry.py").write_text(textwrap.dedent('''
import importlib
from collectors.base_collector import BaseCollector
def load_collectors(config, log):
    modules = {}
    enabled_collectors = config.get("collectors", {}).get("enabled", [])
    for name in enabled_collectors:
        try:
            module = importlib.import_module(f"collectors.{name}_collector")
            for item_name in dir(module):
                item = getattr(module, item_name)
                if isinstance(item, type) and issubclass(item, BaseCollector) and item is not BaseCollector:
                    modules[name] = item(config, log)
                    break
        except ImportError:
            log.error(f"Kon collector module niet laden: '{name}'")
    return modules
'''))

(root / "engine/scoring.py").write_text(textwrap.dedent('''
import math
from .dynamic_parameters import get_dynamic_params

def e_ai_star(params, weights):
    Vp = min(1.0, 0.5 * params.get('V', 0.5) + 0.3 * params.get('Q', 0.4) + 0.2 * min(1.0, params.get('C', 1) / 3.0))
    score = math.sqrt(
        weights.get('w1', 0.3) * (params.get('P', 1.0) * params.get('W_V', 0.7)) +
        weights.get('w2', 0.3) * (params.get('D_A', 0.7) * params.get('D_B', 0.7)) +
        weights.get('w3', 0.2) * (params.get('T', 0.7) * params.get('A', 0.7)) +
        weights.get('w4', 0.15) * Vp +
        weights.get('w5', 0.05) * params.get('M', 0.2)
    )
    return round(score, 4)

def run_scoring(enriched_clusters, config):
    cfg = config.get('scoring', {})
    weights = cfg.get('eai_weights', {})
    base_defaults = cfg.get('eai_defaults', {})
    
    for asset_key, asset_data in enriched_clusters.items():
        enrich = asset_data.get('enrichment', {})
        dynamic_defaults = get_dynamic_params(asset_data, base_defaults)
        
        asset_params = {
            'P': 1.0, 'V': 1.0 if asset_data.get('validated') else 0.5,
            **enrich, **dynamic_defaults
        }
        asset_score = e_ai_star(asset_params, weights)
        asset_data['score'] = asset_score
        
        for finding in asset_data.get('findings', []):
            finding['score'] = asset_score
            
    return enriched_clusters
'''))

(root / "engine/validation.py").write_text(textwrap.dedent('''
import re
import requests
from urllib.parse import urlparse

def _allowed(host, allowlist):
    host=host.lower()
    for e in [x.strip().lower() for x in allowlist if x.strip()]:
        if e.startswith(".") and host.endswith(e): return True
        if host==e: return True
    return False

def validate_assets(asset_clusters, allowlist, log):
    if not allowlist:
        log.warning('Validatie ingeschakeld, maar allowlist is leeg. Alle assets worden als niet-gevalideerd gemarkeerd.')
        for asset_key, asset_data in asset_clusters.items():
            asset_data['validated'] = False
        return asset_clusters
        
    validated_count = 0
    for asset_key, asset_data in asset_clusters.items():
        host_to_validate = urlparse('https://' + asset_key).hostname or asset_key
        
        if not _allowed(host_to_validate, allowlist):
            asset_data['validated'] = False
            continue
            
        try:
            r = requests.head('https://' + asset_key, timeout=7, allow_redirects=True, headers={'User-Agent':'NGBSE-14.0'})
            ok = 200 <= r.status_code < 400
            asset_data['validated'] = ok
            if ok: validated_count += 1
        except requests.RequestException:
            asset_data['validated'] = False
            
    log.info(f"{validated_count} van de {len(asset_clusters)} unieke assets succesvol gevalideerd.")
    return asset_clusters
'''))

(root / "engine/dynamic_parameters.py").write_text(textwrap.dedent('''
def get_dynamic_params(asset_data, base_defaults):
    params = base_defaults.copy()
    enrich = asset_data.get('enrichment', {})
    
    if enrich.get('M', 0.2) >= 0.8:
        params['T'] = max(0.1, params.get('T', 0.7) * 0.5)
        
    base_v = 1.0 if asset_data.get('validated') else 0.5
    if enrich.get('C', 0) >= 3:
        params['V'] = min(1.0, base_v * 1.2)
    else:
        params['V'] = base_v
        
    return params
'''))

(root / "engine/predictive_synthesis.py").write_text(textwrap.dedent('''
import json, math, os
from collections import defaultdict
import google.generativeai as genai

def query_llm(prompt, config, log):
    synthesis_cfg = config.get("predictive_synthesis", {})
    provider = synthesis_cfg.get("llm_provider")
    
    if provider == "google_gemini":
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            msg = "GEMINI_API_KEY niet gevonden. Synthese wordt overgeslagen."
            log.error(msg)
            return msg
            
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(synthesis_cfg.get("llm_model", "gemini-1.5-flash"))
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            log.error("Fout bij aanroepen Gemini API", {"error": str(e)})
            return f"Error: Fout bij communicatie met Gemini API."
            
    log.error(f"Onbekende LLM provider: {provider}")
    return "Error: Onbekende LLM provider geconfigureerd."

def generate_synthesis(scored_assets, config, log):
    synthesis_cfg = config.get("predictive_synthesis", {})
    if not synthesis_cfg.get("enabled", False): return {}
    
    themes_def = synthesis_cfg.get("themes", {})
    for asset_key, asset_data in scored_assets.items():
        asset_themes = set()
        text_to_scan = asset_key.lower() + " " + " ".join(f['seed']['seed'].lower() for f in asset_data['findings'])
        for theme, definition in themes_def.items():
            if any(kw in text_to_scan for kw in definition.get("keywords", [])):
                asset_themes.add(theme)
        asset_data['themes'] = list(asset_themes)
        
    scenarios = defaultdict(lambda: {"assets": []})
    for asset_key, asset_data in scored_assets.items():
        for theme in asset_data.get('themes', []):
            scenarios[theme]['assets'].append(asset_data)
            
    prior = synthesis_cfg.get("prior_probability", 0.1); prior_odds = prior / (1 - prior) if prior < 1 else float('inf')
    output = {}
    
    for theme, data in scenarios.items():
        items = data["assets"]
        if not items: continue
        
        total_score = sum(a['score'] * len(a['findings']) for a in items)
        total_evidence = sum(len(a['findings']) for a in items)
        avg_score = total_score / total_evidence if total_evidence > 0 else 0
        
        likelihood_ratio = 1 + 9 * (max(0, avg_score - 0.5) / 0.5)
        posterior_odds = prior_odds * likelihood_ratio; probability = posterior_odds / (1 + posterior_odds)
        
        prompt = (
            f"Als NL strategic analyst, geef een intelligence brief voor risicothema '{theme}'.\\n"
            f"BLUF: Kritiekste risico in 1 zin.\\n"
            f"RISICO: Details o.b.v. bewijs (gemiddelde E_AI* score {avg_score:.2f}).\\n"
            f"VERTROUWEN: Hoog/Medium/Laag, o.b.v. bewijskwaliteit.\\n"
            f"ACTIE: Concrete, niet-intrusieve vervolgstap voor een analist."
        )
        synthesis_text = query_llm(prompt, config, log)
        
        output[theme] = {
            "future_scenario": themes_def.get(theme, {}).get("future_scenario", "Onbekend"),
            "probability_90_days": f"{probability:.1%}",
            "intelligence_brief": synthesis_text,
            "avg_eai_score": avg_score,
            "asset_count": len(items)
        }
    return dict(sorted(output.items(), key=lambda item: float(item[1]['probability_90_days'][:-1]), reverse=True))
'''))
(root / "engine/reverse_llm.py").write_text(textwrap.dedent('''
from collections import defaultdict

def analyze_blindspots(scored_assets, synthesis):
    blindspots = []
    if not scored_assets:
        return ["Geen assets gevonden. Overweeg meer diverse seeds."]
    total_assets = len(scored_assets)
    ghost_asset_count = sum(1 for data in scored_assets.values() if data.get('enrichment', {}).get('M', 0) >= 0.8)
    if ghost_asset_count > 0 and ghost_asset_count / total_assets > 0.5:
        blindspots.append(f"Strategische Blindspot: {ghost_asset_count}/{total_assets} assets zijn 'ghost assets'. Risico op onbeheerde voetafdruk is hoog.")
    low_corroboration_assets = sum(1 for data in scored_assets.values() if data.get('enrichment', {}).get('C', 1) <= 1 and data.get('enrichment',{}).get('M',1.0) < 0.8)
    if low_corroboration_assets > 0 and low_corroboration_assets / total_assets > 0.6:
         blindspots.append("Evidence Blindspot: De meerderheid van de live assets is slechts door één bron gevonden. Corroboratie is nodig.")
    for theme, data in synthesis.items():
        try:
            p = float(data['probability_90_days'][:-1])
            if p > 75.0:
                blindspots.append(f"Actiegerichte Blindspot: Hoog-risico scenario '{data['future_scenario']}' ({p:.1f}%) vereist onmiddellijke triage.")
        except Exception: continue
    if not blindspots:
        return ["Analyse in balans. Overweeg specialistische seeds."]
    return blindspots
'''))

# --- VERRIJKING, EXPORTS, RAPPORTAGE (VOLLEDIG) ---
(root / "enrich/asset_clustering.py").write_text(textwrap.dedent('''
from urllib.parse import urlparse; from collections import defaultdict
def normalize_asset(finding):
    url = finding.get('url') or finding.get('where');
    if not url: return None
    try:
        p = urlparse(url.lower()); host = p.hostname or ""
        if host.startswith("www."): host = host[4:]
        if any(ext in p.path for ext in ['.yml','.json','.config']): return f"{host}{p.path}"
        return host
    except: return None
def cluster_by_asset(findings):
    clusters = defaultdict(lambda: {"findings": [], "urls": set()})
    for f in findings:
        asset_key = normalize_asset(f)
        if asset_key and f.get('url') not in clusters[asset_key]['urls']:
            clusters[asset_key]['findings'].append(f)
            clusters[asset_key]['urls'].add(f.get('url'))
    return dict(clusters)
'''))
(root / "enrich/enrichment.py").write_text(textwrap.dedent('''
from collections import defaultdict
def run_enrichment(asset_clusters, config):
    quality_weights = config.get('scoring', {}).get('quality_weights', {})
    for asset_key, asset_data in asset_clusters.items():
        sources = {f['source'] for f in asset_data['findings']}
        live_sources = sources - {'wayback'}
        C = len(live_sources)
        M = 1.0 if 'wayback' in sources and not live_sources else 0.6 if 'wayback' in sources else 0.2
        Q = sum(quality_weights.get(f['source'], 0.4) for f in asset_data['findings']) / len(asset_data['findings']) if asset_data['findings'] else 0.4
        asset_data['enrichment'] = {'M': M, 'C': C, 'Q': round(Q, 2)}
    return asset_clusters
'''))
(root / "export/csv_export.py").write_text(textwrap.dedent('''
import csv
def write_csv(findings, path):
    if not findings: return
    flat_findings = []
    for f in findings:
        ff = dict(f); enrich = ff.pop('enrichment', {}); ff.update(enrich); ff.pop('themes', None)
        seed_data = ff.pop('seed', {}); ff['seed_query'] = seed_data.get('seed', '')
        flat_findings.append(ff)
    all_keys = sorted({k for item in flat_findings for k in item.keys()})
    with open(path, 'w', newline='', encoding='utf-8') as h:
        w = csv.DictWriter(h, fieldnames=all_keys, extrasaction='ignore'); w.writeheader(); w.writerows(flat_findings)
'''))
(root / "export/stix_exporter.py").write_text(textwrap.dedent('''
import json, uuid, datetime
from stix2 import TLP_WHITE, Indicator, Bundle, Report, Identity
def export_to_stix(findings, synthesis, author_name, out_path):
    stix_objects, report_refs = [], []
    identity = Identity(name=author_name, identity_class="organization")
    stix_objects.append(identity)
    for f in findings:
        url = f.get("url") or f.get("where", "")
        pattern = f"[url:value = '{url}']" if f.get("url") else f"[{f.get('source')}:value = '{f.get('where')}']"
        if not (f.get("url") or f.get("where")): continue
        indicator = Indicator(name=f"OSINT Finding: {url}", pattern_type="stix", pattern=pattern,
            valid_from=f.get("ts") or datetime.datetime.utcnow(), created_by_ref=identity.id,
            labels=["osint-finding"], confidence=int(f.get("score", 0.5) * 100))
        stix_objects.append(indicator); report_refs.append(indicator.id)
    synthesis_desc = "\\n".join([f"{data['future_scenario']} (Prob: {data['probability_90_days']})" for theme, data in synthesis.items()])
    stix_report = Report(name=f"NGBSE Final Synthesis ({datetime.datetime.now().strftime('%Y-%m-%d')})",
        description=synthesis_desc, object_refs=report_refs, created_by_ref=identity.id, report_types=["threat-report"])
    stix_objects.append(stix_report)
    bundle = Bundle(stix_objects, allow_custom=True)
    with open(out_path, "w", encoding="utf-8") as f: f.write(bundle.serialize(pretty=True))
'''))
(root / "report/make_report.py").write_text(textwrap.dedent('''
import json
from pathlib import Path
from docx import Document
def build_report(scored_assets, synthesis, blindspots):
    out_path = Path("out/NGBSE_FINAL_Report.docx"); doc = Document()
    doc.add_heading("NGBSE-FINAL: Executive Briefing", 0)
    doc.add_heading("Voorspellende Risicoscenario's", 1)
    if not synthesis: doc.add_paragraph("Geen risicoscenario's gedetecteerd.")
    for theme, data in synthesis.items():
        doc.add_paragraph(data['future_scenario'], style='Heading 2')
        p = doc.add_paragraph(); p.add_run('Waarschijnlijkheid (90 dagen): ').bold = True
        p.add_run(f'{data["probability_90_days"]}\\n'); p.add_run('Intelligence Brief: ').bold = True
        p.add_run(data['intelligence_brief'])
    doc.add_heading("Gecombineerde Blindspots & Aanbevelingen", 1)
    for b in blindspots: doc.add_paragraph(b, style="List Bullet")
    doc.add_heading('Top Risico-Assets', 1); table = doc.add_table(1, 4); table.style = 'Table Grid'
    hdr = table.rows[0].cells; hdr[0].text = 'E_AI*'; hdr[1].text = 'Asset'; hdr[2].text = 'Info (M,C,Q)'; hdr[3].text = 'Bewijs'
    sorted_assets = sorted(scored_assets.values(), key=lambda x: x.get('score', 0), reverse=True)
    for asset in sorted_assets[:20]:
        r = table.add_row().cells; r[0].text = f'{asset.get("score",0):.4f}'
        r[1].text = list(asset['urls'])[0] if asset.get('urls') else 'N/A'
        e = asset.get('enrichment', {}); r[2].text = f"M={e.get('M')} C={e.get('C')} Q={e.get('Q'):.2f}"
        r[3].text = f'{len(asset["findings"])} indicatoren'
    doc.save(out_path); return out_path
'''))

# --- OPERATIONELE COLLECTORS (VOLLEDIG & OBJECTGEORIËNTEERD) ---
(root / "collectors/base_collector.py").write_text("import os, requests, time\nfrom .util import coc_sha256\nclass BaseCollector:\n    def __init__(self,c,l): self.config=c; self.log=l; self.source_name='unknown'; self.api_key=None; self.requires_api_key=False; self.api_key_env_var=None; self.session=requests.Session(); self.session.headers.update({'User-Agent':'NGBSE-FINAL'}); self._load_api_key()\n    def _load_api_key(self): \n        if self.requires_api_key and self.api_key_env_var: self.api_key=os.getenv(self.api_key_env_var)\n    def collect(self,s,c,l): \n        is_c=c.get('orchestrator',{}).get('clearnet_only_mode',False)\n        if self.requires_api_key and (not self.api_key or is_c): self.log.warning(f'{self.source_name.capitalize()}: API-sleutel vereist.', {'seed': s.get('seed')}); return []\n        return self.run(s)\n    def run(self,s): raise NotImplementedError()")
(root / "collectors/util.py").write_text("import hashlib\ndef coc_sha256(m:str)->str: return hashlib.sha256(m.encode('utf-8','ignore')).hexdigest()")
(root / "collectors/urlscan_collector.py").write_text("from .base_collector import BaseCollector\nfrom .util import coc_sha256\nimport requests\nclass UrlscanCollector(BaseCollector):\n    def __init__(self,c,l): super().__init__(c,l); self.source_name='urlscan'; self.requires_api_key=True; self.api_key_env_var='URLSCAN_API_KEY'; self._load_api_key(); \n        if self.api_key: self.session.headers.update({'API-Key':self.api_key})\n    def run(self,s): \n        q=s.get('seed'); \n        try: r=self.session.get('https://urlscan.io/api/v1/search/',params={'q':q,'size':50},timeout=20); r.raise_for_status()\n        except Exception as e: self.log.error(f'{self.source_name} Fout',{'error':str(e),'seed':q}); return\n        for i in r.json().get('results',[]): u=i.get('page',{}).get('url'); ts=i.get('indexedAt'); rid=i.get('_id'); yield {'source':'urlscan','seed':s,'url':u,'ts':ts,'coc_sha256':coc_sha256(f'{rid}|{u}|{ts}')}")
(root / "collectors/github_collector.py").write_text("from .base_collector import BaseCollector\nfrom .util import coc_sha256\nimport requests\nfrom bs4 import BeautifulSoup\nfrom urllib.parse import quote_plus\nclass GithubCollector(BaseCollector):\n    def __init__(self,c,l): super().__init__(c,l); self.source_name='github'; self.api_key_env_var='GITHUB_TOKEN'; self._load_api_key()\n    def collect(self,s,c,l): return self._run_api(s) if self.api_key and not c.get('orchestrator',{}).get('clearnet_only_mode') else self._run_scrape(s)\n    def _run_api(self,s): \n        q=s.get('seed'); self.session.headers.update({'Authorization':f'Bearer {self.api_key}'})\n        try: r=self.session.get('https://api.github.com/search/code',params={'q':q,'per_page':50},timeout=20); r.raise_for_status()\n        except Exception as e: self.log.error(f'{self.source_name} API Fout',{'error':str(e),'seed':q}); return\n        for i in r.json().get('items',[]): u,repo=i.get('html_url'),i.get('repository',{}); yield {'source':'github','seed':s,'url':u,'ts':repo.get('updated_at'),'coc_sha256':coc_sha256(f'{u}|{repo.get(\"full_name\")}')}\n    def _run_scrape(self,s):\n        q=s.get('seed'); self.log.warning('GitHub draait in clearnet-scrape-modus.')\n        try: r=self.session.get(f'https://github.com/search?type=code&q={quote_plus(q)}',timeout=20); r.raise_for_status(); soup=BeautifulSoup(r.text,'html.parser')\n        except Exception as e: self.log.error(f'{self.source_name} Scrape Fout',{'error':str(e),'seed':q}); return\n        for a in soup.select('div.search-title a'): u='https://github.com'+a['href']; yield {'source':'github','seed':s,'url':u,'ts':None,'coc_sha256':coc_sha256(u)}")
(root / "collectors/shodan_collector.py").write_text("from .base_collector import BaseCollector\nfrom .util import coc_sha256\nimport requests\nclass ShodanCollector(BaseCollector):\n    def __init__(self,c,l): super().__init__(c,l); self.source_name='shodan'; self.requires_api_key=True; self.api_key_env_var='SHODAN_API_KEY'; self._load_api_key()\n    def run(self,s): \n        q=s.get('seed'); \n        try: r=self.session.get('https://api.shodan.io/shodan/host/search', params={'key':self.api_key,'query':q},timeout=20); r.raise_for_status()\n        except Exception as e: self.log.error(f'{self.source_name} Fout',{'error':str(e),'seed':q}); return\n        for i in r.json().get('matches',[]): ip=i.get('ip_str'); ts=i.get('timestamp'); yield {'source':'shodan','seed':s,'where':ip,'url':f'https://www.shodan.io/host/{ip}','ts':ts,'coc_sha256':coc_sha256(f'{ip}|{ts}')}")
(root / "collectors/censys_collector.py").write_text("from .base_collector import BaseCollector\nfrom .util import coc_sha256\nimport os, requests\nclass CensysCollector(BaseCollector):\n    def __init__(self,c,l): super().__init__(c,l); self.source_name='censys'; self.requires_api_key=True; self.api_id=os.getenv('CENSYS_API_ID'); self.api_secret=os.getenv('CENSYS_API_SECRET')\n    def collect(self,s,c,l): \n        if not (self.api_id and self.api_secret) or c.get('orchestrator',{}).get('clearnet_only_mode'): self.log.warning('Censys: API ID/Secret vereist.'); return []\n        return self.run(s)\n    def run(self,s): \n        q=s.get('seed')\n        try: r=self.session.post('https://search.censys.io/api/v2/hosts/search',auth=(self.api_id,self.api_secret),json={'q':q,'per_page':50},timeout=20); r.raise_for_status()\n        except Exception as e: self.log.error(f'{self.source_name} Fout',{'error':str(e),'seed':q}); return\n        for i in r.json().get('result',{}).get('hits',[]): ip=i.get('ip'); ts=i.get('last_updated_at'); yield {'source':'censys','seed':s,'where':ip,'url':f'https://search.censys.io/hosts/{ip}','ts':ts,'coc_sha256':coc_sha256(f'{ip}|{ts}')}")
(root / "collectors/leakix_collector.py").write_text("from .base_collector import BaseCollector\nfrom .util import coc_sha256\nimport requests\nclass LeakixCollector(BaseCollector):\n    def __init__(self,c,l): super().__init__(c,l); self.source_name='leakix'; self.requires_api_key=True; self.api_key_env_var='LEAKIX_API_KEY'; self._load_api_key()\n        if self.api_key: self.session.headers.update({'Api-Key':self.api_key})\n    def run(self,s): \n        q=s.get('seed')\n        try: r=self.session.get('https://leakix.net/search',params={'q':q,'size':50},timeout=20); r.raise_for_status()\n        except Exception as e: self.log.error(f'{self.source_name} Fout',{'error':str(e),'seed':q}); return\n        for i in r.json(): u,ts=i.get('link'),i.get('time'); yield {'source':'leakix','seed':s,'url':u,'where':i.get('ip'),'ts':ts,'coc_sha256':coc_sha256(f'{u}|{ts}')}")
(root / "collectors/wayback_collector.py").write_text("from .base_collector import BaseCollector\nfrom .util import coc_sha256\nimport requests\nclass WaybackCollector(BaseCollector):\n    def __init__(self,c,l): super().__init__(c,l); self.source_name='wayback'\n    def run(self,s): \n        q=s.get('seed')\n        try: r=self.session.get('http://web.archive.org/cdx/search/cdx',params={'url':f'*{q}*','output':'json','limit':100},timeout=20); r.raise_for_status()\n        except Exception as e: self.log.error(f'{self.source_name} Fout',{'error':str(e),'seed':q}); return\n        for _,ts,o,_,_,d,_ in r.json()[1:]: u,t_iso=f'https://web.archive.org/web/{ts}/{o}',f'{ts[:4]}-{ts[4:6]}-{ts[6:8]}T{ts[8:10]}:{ts[10:12]}:{ts[12:14]}Z'; yield {'source':'wayback','seed':s,'url':u,'ts':t_iso,'coc_sha256':coc_sha256(f'{u}|{d}')}")

# --- RUN SCRIPT ---
(root / "run.sh").write_text(textwrap.dedent('''
#!/bin/bash
set -e
echo ">>> NGBSE-FINAL: De Finale Engine starten..."
if [ ! -f .env ]; then
    cp .env.example .env
    echo "BELANGRIJK: Vul API-sleutels in .env voor de beste resultaten."
fi
pip install -q -r requirements.txt
python -m engine.orchestrator "$@"
'''))
os.chmod(root / "run.sh", 0o755)
(root / "README.md").write_text(textwrap.dedent('''
# NGBSE-FINAL: De Finale, Complete Engine
Dit is de finale, zuivere, online-gerichte NGBSE, die de meest geavanceerde concepten van alle voorgaande versies combineert.

## Vereisten
- Python 3.9+
- Een lokaal draaiende [Ollama](https://ollama.com/) instance met het `phi3:mini` model. (`ollama serve` & `ollama pull phi3:mini`)

## Gebruik
1. `cp .env.example .env` (en vul API-sleutels in).
2. `./run.sh`
'''))

print("\\n>>> SETUP VOLTOOID. NGBSE-FINAL is gebouwd. <<<")
print(f"1. Navigeer naar de map: cd {root.name}")
print("2. Vul uw API-sleutels in '.env'.")
print("3. Zorg dat Ollama lokaal draait (`ollama serve`).")
print("4. Start de engine met: ./run.sh")