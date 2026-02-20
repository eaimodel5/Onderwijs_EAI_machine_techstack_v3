try:
    from docx import Document
    from docx.shared import Pt
    _DOCX = True
except Exception:
    _DOCX = False
    Document = None
    Pt = None

from typing import List, Dict, Any
from datetime import datetime

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11) if Pt else None
    return p

def write_report(path: str, brief: Dict[str, Any], blindspots: Dict[str, Any], forecast: Dict[str, Any], scenarios: Dict[str, Any] | None = None):
    if not _DOCX:
        # Fallback: write markdown next to expected path
        md_path = path.rsplit('.', 1)[0] + ".md"
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write("# NGBSE 17.1 – Intelligence Brief\n")
            f.write(f"Generated: {datetime.utcnow().isoformat()}Z\n\n")
            f.write("## BLUF\n" + brief.get("BLUF", "") + "\n\n")
            f.write("## Top Risks\n")
            for r in brief.get("TopRisks", []):
                f.write(f"- {r['asset']}: risk {r['risk']}\n")
            f.write("\n## Actions\n")
            for a in brief.get("Actions", []):
                f.write(f"- {a}\n")
            f.write("\n## Blindspot Analysis\n")
            f.write(f"Coverage Ratios: {blindspots['coverage'].get('ratios', {})}\n")
            f.write(f"Coverage Imbalance: {blindspots['coverage'].get('imbalance', 0.0)}\n")
            f.write(f"Recency Spike: {blindspots['recency'].get('spike', False)} / {blindspots['recency'].get('detail', {})}\n")
            f.write(f"Confidence Gap (risky assets): {blindspots['confidence'].get('risky_assets', [])}\n")
            f.write("\n## Forecast\n")
            if forecast:
                for asset, fvals in forecast.items():
                    f.write(f"- {asset}: last={round(fvals['last'], 3)} ewma={round(fvals['ewma'], 3)} → projected_next={round(fvals['projected_next'], 3)}\n")
            else:
                f.write("No historical series available yet.\n")
            if scenarios:
                f.write("\n## Scenarios (90 dagen)\n")
                for theme, data in scenarios.items():
                    f.write(f"- {theme}: p≈{round(100*data.get('probability_90_days',0.0),1)}% | {data.get('semantic_summary','').strip()}\n")
            f.write("\n## Analyst Review & Sign-off\n")
            f.write("Analyst Name: ________________________\nDate of Review: ______________________\nConclusion: [Agree] [Disagree]\nSignature: ___________________________\n")
        return

    doc = Document()
    add_heading(doc, "NGBSE 17.1 – Intelligence Brief", 0)
    add_paragraph(doc, f"Generated: {datetime.utcnow().isoformat()}Z")

    add_heading(doc, "BLUF", 1)
    add_paragraph(doc, brief.get("BLUF", ""))

    add_heading(doc, "Top Risks", 1)
    for r in brief.get("TopRisks", []):
        add_paragraph(doc, f"- {r['asset']}: risk {r['risk']}")

    add_heading(doc, "Actions", 1)
    for a in brief.get("Actions", []):
        add_paragraph(doc, f"- {a}")

    add_heading(doc, "Blindspot Analysis", 1)
    add_paragraph(doc, f"Coverage Ratios: {blindspots['coverage'].get('ratios', {})}")
    add_paragraph(doc, f"Coverage Imbalance: {blindspots['coverage'].get('imbalance', 0.0)}")
    add_paragraph(doc, f"Recency Spike: {blindspots['recency'].get('spike', False)} / {blindspots['recency'].get('detail', {})}")
    add_paragraph(doc, f"Confidence Gap (risky assets): {blindspots['confidence'].get('risky_assets', [])}")

    add_heading(doc, "Forecast & Scenarios", 1)
    if forecast:
        for asset, fvals in forecast.items():
            add_paragraph(doc, f"- {asset}: last={round(fvals['last'], 3)} ewma={round(fvals['ewma'], 3)} → projected_next={round(fvals['projected_next'], 3)}")
    else:
        add_paragraph(doc, "No historical series available yet.")
    if scenarios:
        add_heading(doc, "Scenarios (90 dagen)", 2)
        for theme, data in scenarios.items():
            prob = round(100*data.get("probability_90_days", 0.0), 1)
            summary = (data.get("semantic_summary") or "").strip()
            add_paragraph(doc, f"- {theme}: p≈{prob}% {('— ' + summary) if summary else ''}")

    add_heading(doc, "Analyst Review & Sign-off", 1)
    add_paragraph(doc, "Analyst Name: ________________________")
    add_paragraph(doc, "Date of Review: ______________________")
    add_paragraph(doc, "Conclusion: [Agree] [Disagree]")
    add_paragraph(doc, "Signature: ___________________________")

    doc.save(path)
